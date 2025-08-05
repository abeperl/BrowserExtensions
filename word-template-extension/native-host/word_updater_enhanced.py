#!/usr/bin/env python3
"""
Enhanced Word Template Updater - Native Messaging Host
Receives data from browser extension and updates Word document templates.
Includes improved error handling, logging, and diagnostics.
"""

import sys
import json
import struct
import logging
import os
import subprocess
import traceback
from datetime import datetime
from pathlib import Path
from typing import Dict, Any, Optional, List

try:
    from docx import Document
    from docx.shared import Inches
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# Configure logging with rotation
log_dir = Path.home() / "AppData" / "Local" / "WordTemplateExtension"
log_dir.mkdir(parents=True, exist_ok=True)
log_file = log_dir / "word_updater.log"

# Set up logging with more detailed format
logging.basicConfig(
    level=logging.DEBUG,
    format='%(asctime)s - %(levelname)s - [%(funcName)s:%(lineno)d] %(message)s',
    handlers=[
        logging.FileHandler(log_file),
        logging.StreamHandler(sys.stderr)
    ]
)
logger = logging.getLogger(__name__)

class NativeMessagingError(Exception):
    """Custom exception for native messaging errors."""
    pass

class WordTemplateUpdaterEnhanced:
    """Enhanced version with better error handling and diagnostics."""
    
    def __init__(self):
        logger.info("Initializing WordTemplateUpdaterEnhanced")
        self.config_dir = Path.home() / "AppData" / "Local" / "WordTemplateExtension"
        self.config_dir.mkdir(parents=True, exist_ok=True)
        self.config_file = self.config_dir / "config.json"
        
        # Check dependencies first
        self.check_dependencies()
        self.load_config()
        
        logger.info("WordTemplateUpdaterEnhanced initialized successfully")
    
    def check_dependencies(self):
        """Check if all required dependencies are available."""
        if not DOCX_AVAILABLE:
            error_msg = "python-docx library not installed. Run: pip install python-docx"
            logger.error(error_msg)
            self.send_error_response("dependency_error", error_msg)
            sys.exit(1)
        
        logger.info("All dependencies available")
    
    def load_config(self):
        """Load configuration settings with validation."""
        default_config = {
            "template_path": str(Path.home() / "Documents" / "Templates"),
            "output_path": str(Path.home() / "Documents" / "Generated"),
            "auto_open": True,
            "default_template": "template.docx",
            "max_file_size_mb": 50,
            "allowed_extensions": [".docx", ".docm"],
            "log_level": "INFO"
        }
        
        try:
            if self.config_file.exists():
                with open(self.config_file, 'r', encoding='utf-8') as f:
                    user_config = json.load(f)
                self.config = {**default_config, **user_config}
                logger.info("Configuration loaded from file")
            else:
                self.config = default_config
                self.save_config()
                logger.info("Default configuration created")
                
            # Validate and create directories
            self.validate_config()
            
        except Exception as e:
            logger.error(f"Error loading config: {e}")
            self.config = default_config
            self.validate_config()
    
    def validate_config(self):
        """Validate configuration and create necessary directories."""
        try:
            # Create template directory
            template_path = Path(self.config["template_path"])
            template_path.mkdir(parents=True, exist_ok=True)
            
            # Create output directory
            output_path = Path(self.config["output_path"])
            output_path.mkdir(parents=True, exist_ok=True)
            
            logger.info(f"Template directory: {template_path}")
            logger.info(f"Output directory: {output_path}")
            
        except Exception as e:
            logger.error(f"Error validating config: {e}")
            raise NativeMessagingError(f"Configuration validation failed: {e}")
    
    def save_config(self):
        """Save configuration settings."""
        try:
            with open(self.config_file, 'w', encoding='utf-8') as f:
                json.dump(self.config, f, indent=2)
            logger.debug("Configuration saved")
        except Exception as e:
            logger.error(f"Error saving config: {e}")
    
    def read_message(self) -> Optional[Dict[str, Any]]:
        """Read a message from stdin using Chrome native messaging format."""
        try:
            # Read the message length (4 bytes, little-endian)
            raw_length = sys.stdin.buffer.read(4)
            if not raw_length:
                logger.debug("No message length received, assuming shutdown")
                return None
            
            if len(raw_length) != 4:
                raise NativeMessagingError(f"Invalid message length header: {len(raw_length)} bytes")
            
            message_length = struct.unpack('<I', raw_length)[0]
            logger.debug(f"Message length: {message_length}")
            
            if message_length > 1024 * 1024:  # 1MB limit
                raise NativeMessagingError(f"Message too large: {message_length} bytes")
            
            # Read the message
            message_data = sys.stdin.buffer.read(message_length)
            if len(message_data) != message_length:
                raise NativeMessagingError(f"Incomplete message: expected {message_length}, got {len(message_data)}")
            
            message_str = message_data.decode('utf-8')
            message = json.loads(message_str)
            
            logger.debug(f"Received message: {json.dumps(message, indent=2)}")
            return message
            
        except json.JSONDecodeError as e:
            logger.error(f"Invalid JSON in message: {e}")
            self.send_error_response("invalid_json", f"Invalid JSON: {e}")
            return None
        except Exception as e:
            logger.error(f"Error reading message: {e}")
            self.send_error_response("read_error", f"Message read error: {e}")
            return None
    
    def send_message(self, message: Dict[str, Any]):
        """Send a message to stdout using Chrome native messaging format."""
        try:
            encoded_message = json.dumps(message, ensure_ascii=False).encode('utf-8')
            message_length = len(encoded_message)
            
            logger.debug(f"Sending message ({message_length} bytes): {json.dumps(message, indent=2)}")
            
            # Send message length (4 bytes, little-endian)
            sys.stdout.buffer.write(struct.pack('<I', message_length))
            
            # Send message
            sys.stdout.buffer.write(encoded_message)
            sys.stdout.buffer.flush()
            
        except Exception as e:
            logger.error(f"Error sending message: {e}")
            # Can't send error response if send_message itself fails
    
    def send_error_response(self, error_type: str, error_message: str, **kwargs):
        """Send standardized error response."""
        response = {
            "success": False,
            "error": {
                "type": error_type,
                "message": error_message,
                "timestamp": datetime.now().isoformat(),
                **kwargs
            }
        }
        self.send_message(response)
    
    def send_success_response(self, data: Dict[str, Any] = None, **kwargs):
        """Send standardized success response."""
        response = {
            "success": True,
            "timestamp": datetime.now().isoformat(),
            **(data or {}),
            **kwargs
        }
        self.send_message(response)
    
    def handle_ping(self, message: Dict[str, Any]) -> Dict[str, Any]:
        """Handle ping requests for connectivity testing."""
        logger.info("Ping request received")
        return {
            "action": "pong",
            "version": "2.0.0",
            "status": "ready",
            "config": {
                "template_path": self.config["template_path"],
                "output_path": self.config["output_path"],
                "docx_available": DOCX_AVAILABLE
            }
        }
    
    def handle_list_templates(self, message: Dict[str, Any]) -> Dict[str, Any]:
        """List available templates."""
        try:
            template_path = Path(self.config["template_path"])
            templates = []
            
            if template_path.exists():
                for ext in self.config["allowed_extensions"]:
                    for template_file in template_path.glob(f"*{ext}"):
                        templates.append({
                            "name": template_file.name,
                            "path": str(template_file),
                            "size": template_file.stat().st_size,
                            "modified": datetime.fromtimestamp(template_file.stat().st_mtime).isoformat()
                        })
            
            logger.info(f"Found {len(templates)} templates")
            return {
                "action": "template_list",
                "templates": templates,
                "template_path": str(template_path)
            }
            
        except Exception as e:
            logger.error(f"Error listing templates: {e}")
            raise NativeMessagingError(f"Failed to list templates: {e}")
    
    def handle_update_template(self, message: Dict[str, Any]) -> Dict[str, Any]:
        """Update a Word template with provided data."""
        try:
            data = message.get("data", {})
            template_name = data.get("template")
            extracted_data = data.get("extractedData", {})
            
            if not template_name:
                raise NativeMessagingError("No template specified")
            
            template_path = Path(self.config["template_path"]) / template_name
            if not template_path.exists():
                raise NativeMessagingError(f"Template not found: {template_name}")
            
            # Check file size
            template_size = template_path.stat().st_size
            max_size = self.config["max_file_size_mb"] * 1024 * 1024
            if template_size > max_size:
                raise NativeMessagingError(f"Template too large: {template_size / 1024 / 1024:.1f}MB")
            
            logger.info(f"Processing template: {template_name}")
            logger.debug(f"Extracted data: {extracted_data}")
            
            # Process the template
            output_path = self.process_template(template_path, extracted_data)
            
            return {
                "action": "template_updated",
                "template": template_name,
                "output_path": str(output_path),
                "data_processed": len(extracted_data)
            }
            
        except Exception as e:
            logger.error(f"Error updating template: {e}")
            raise NativeMessagingError(f"Template update failed: {e}")
    
    def process_template(self, template_path: Path, data: Dict[str, Any]) -> Path:
        """Process a template with provided data."""
        try:
            # Load template
            doc = Document(str(template_path))
            
            # Generate output filename
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            base_name = template_path.stem
            output_name = f"{base_name}_{timestamp}.docx"
            output_path = Path(self.config["output_path"]) / output_name
            
            # Replace placeholders
            replacements_made = 0
            
            # Process paragraphs
            for paragraph in doc.paragraphs:
                original_text = paragraph.text
                new_text = self.replace_placeholders(original_text, data)
                if new_text != original_text:
                    paragraph.text = new_text
                    replacements_made += 1
            
            # Process tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        original_text = cell.text
                        new_text = self.replace_placeholders(original_text, data)
                        if new_text != original_text:
                            cell.text = new_text
                            replacements_made += 1
            
            # Save document
            doc.save(str(output_path))
            
            logger.info(f"Template processed: {replacements_made} replacements made")
            logger.info(f"Output saved to: {output_path}")
            
            # Auto-open if configured
            if self.config.get("auto_open", False):
                try:
                    os.startfile(str(output_path))
                    logger.info("Document opened automatically")
                except Exception as e:
                    logger.warning(f"Could not auto-open document: {e}")
            
            return output_path
            
        except Exception as e:
            logger.error(f"Error processing template: {e}")
            logger.error(f"Traceback: {traceback.format_exc()}")
            raise
    
    def replace_placeholders(self, text: str, data: Dict[str, Any]) -> str:
        """Replace placeholders in text with data values."""
        if not text or not data:
            return text
        
        result = text
        for key, value in data.items():
            placeholder = f"{{{{{key}}}}}"
            if placeholder in result:
                # Convert value to string, handling different data types
                str_value = self.format_value(value)
                result = result.replace(placeholder, str_value)
                logger.debug(f"Replaced {placeholder} with '{str_value}'")
        
        return result
    
    def format_value(self, value: Any) -> str:
        """Format a value for insertion into document."""
        if value is None:
            return ""
        elif isinstance(value, (list, tuple)):
            return ", ".join(str(v) for v in value)
        elif isinstance(value, dict):
            return json.dumps(value)
        elif isinstance(value, datetime):
            return value.strftime("%Y-%m-%d %H:%M:%S")
        else:
            return str(value)
    
    def handle_get_config(self, message: Dict[str, Any]) -> Dict[str, Any]:
        """Get current configuration."""
        return {
            "action": "config",
            "config": self.config.copy()
        }
    
    def handle_update_config(self, message: Dict[str, Any]) -> Dict[str, Any]:
        """Update configuration."""
        try:
            new_config = message.get("config", {})
            
            # Validate new config
            for key, value in new_config.items():
                if key in self.config:
                    self.config[key] = value
            
            self.save_config()
            self.validate_config()
            
            logger.info("Configuration updated")
            return {
                "action": "config_updated",
                "config": self.config.copy()
            }
            
        except Exception as e:
            logger.error(f"Error updating config: {e}")
            raise NativeMessagingError(f"Configuration update failed: {e}")
    
    def process_message(self, message: Dict[str, Any]):
        """Process incoming messages and route to appropriate handlers."""
        try:
            action = message.get("action", "unknown")
            logger.info(f"Processing action: {action}")
            
            handlers = {
                "ping": self.handle_ping,
                "list_templates": self.handle_list_templates,
                "update_template": self.handle_update_template,
                "get_config": self.handle_get_config,
                "update_config": self.handle_update_config
            }
            
            if action in handlers:
                response_data = handlers[action](message)
                self.send_success_response(response_data)
            else:
                error_msg = f"Unknown action: {action}"
                logger.error(error_msg)
                self.send_error_response("unknown_action", error_msg, available_actions=list(handlers.keys()))
                
        except NativeMessagingError as e:
            logger.error(f"Native messaging error: {e}")
            self.send_error_response("native_messaging_error", str(e))
        except Exception as e:
            logger.error(f"Unexpected error processing message: {e}")
            logger.error(f"Traceback: {traceback.format_exc()}")
            self.send_error_response("internal_error", f"Internal error: {e}")
    
    def run(self):
        """Main message processing loop."""
        logger.info("Starting WordTemplateUpdaterEnhanced main loop")
        
        try:
            while True:
                message = self.read_message()
                if message is None:
                    logger.info("No message received, shutting down")
                    break
                
                self.process_message(message)
                
        except KeyboardInterrupt:
            logger.info("Received keyboard interrupt, shutting down")
        except Exception as e:
            logger.error(f"Fatal error in main loop: {e}")
            logger.error(f"Traceback: {traceback.format_exc()}")
            self.send_error_response("fatal_error", f"Fatal error: {e}")
        finally:
            logger.info("WordTemplateUpdaterEnhanced shutting down")

def main():
    """Main entry point."""
    try:
        updater = WordTemplateUpdaterEnhanced()
        updater.run()
    except Exception as e:
        logger.error(f"Failed to start: {e}")
        logger.error(f"Traceback: {traceback.format_exc()}")
        sys.exit(1)

if __name__ == "__main__":
    main()