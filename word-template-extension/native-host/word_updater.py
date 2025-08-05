#!/usr/bin/env python3
"""
Word Template Updater - Native Messaging Host
Receives data from browser extension and updates Word document templates.
"""

import sys
import json
import struct
import logging
import os
import subprocess
from datetime import datetime
from pathlib import Path
from typing import Dict, Any, Optional

try:
    from docx import Document
    from docx.shared import Inches
except ImportError:
    print("Error: python-docx not installed. Run: pip install python-docx", file=sys.stderr)
    sys.exit(1)

# Configure logging
log_dir = Path.home() / "AppData" / "Local" / "WordTemplateExtension"
log_dir.mkdir(parents=True, exist_ok=True)
log_file = log_dir / "word_updater.log"

logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler(log_file),
        logging.StreamHandler(sys.stderr)
    ]
)
logger = logging.getLogger(__name__)

class WordTemplateUpdater:
    """Handles Word document template updates via native messaging."""
    
    def __init__(self):
        self.config_dir = Path.home() / "AppData" / "Local" / "WordTemplateExtension"
        self.config_dir.mkdir(parents=True, exist_ok=True)
        self.config_file = self.config_dir / "config.json"
        self.load_config()
    
    def load_config(self):
        """Load configuration settings."""
        default_config = {
            "template_path": str(Path.home() / "Documents" / "Templates"),
            "output_path": str(Path.home() / "Documents" / "Generated"),
            "auto_open": True,
            "default_template": "template.docx"
        }
        
        try:
            if self.config_file.exists():
                with open(self.config_file, 'r') as f:
                    self.config = {**default_config, **json.load(f)}
            else:
                self.config = default_config
                self.save_config()
        except Exception as e:
            logger.error(f"Error loading config: {e}")
            self.config = default_config
    
    def save_config(self):
        """Save configuration settings."""
        try:
            with open(self.config_file, 'w') as f:
                json.dump(self.config, f, indent=2)
        except Exception as e:
            logger.error(f"Error saving config: {e}")
    
    def read_message(self) -> Optional[Dict[str, Any]]:
        """Read a message from stdin using Chrome native messaging format."""
        try:
            # Read the message length (4 bytes, little-endian)
            raw_length = sys.stdin.buffer.read(4)
            if not raw_length:
                return None
            
            message_length = struct.unpack('<I', raw_length)[0]
            
            # Read the message
            message = sys.stdin.buffer.read(message_length).decode('utf-8')
            return json.loads(message)
        
        except Exception as e:
            logger.error(f"Error reading message: {e}")
            return None
    
    def send_message(self, message: Dict[str, Any]):
        """Send a message to stdout using Chrome native messaging format."""
        try:
            encoded_message = json.dumps(message).encode('utf-8')
            message_length = len(encoded_message)
            
            # Send message length (4 bytes, little-endian)
            sys.stdout.buffer.write(struct.pack('<I', message_length))
            # Send message
            sys.stdout.buffer.write(encoded_message)
            sys.stdout.buffer.flush()
            
        except Exception as e:
            logger.error(f"Error sending message: {e}")
    
    def replace_placeholders(self, doc: Document, data: Dict[str, Any]):
        """Replace placeholders in the document with actual data using dynamic mappings."""
        logger.info(f"Data received for replacement: {data}")  # Debug logging
        
        # Get settings including field mappings
        settings = data.get('settings', {})
        field_mappings = settings.get('fieldMappings', [])
        array_handling = settings.get('arrayHandling', 'first')
        text_transform = settings.get('textTransform', 'none')
        preserve_empty = settings.get('preserveEmptyPlaceholders', False)
        
        replacements = {}
        
        # Process dynamic mappings
        for mapping in field_mappings:
            source_field = mapping.get('sourceField', '')
            placeholder = mapping.get('placeholder', '')
            transform = mapping.get('transform', 'none')
            
            if not source_field or not placeholder:
                continue
                
            # Get value from data using dynamic field path
            value = self.get_field_value(data, source_field, array_handling)
            
            # Apply transformations
            if value is not None and value != '':
                value = self.apply_text_transform(str(value), transform if transform != 'none' else text_transform)
            
            # Add to replacements
            if value is not None and (value != '' or preserve_empty):
                replacements[f'{{{{{placeholder}}}}}'] = str(value) if value != '' else f'{{{{{placeholder}}}}}'
        
        # Add default/legacy replacements if no custom mappings exist
        if not field_mappings:
            replacements.update({
                '{{TITLE}}': data.get('title', data.get('pageTitle', '')),
                '{{DATE}}': data.get('date', data.get('extractionDate', datetime.now().strftime('%Y-%m-%d'))),
                '{{AMOUNT}}': data.get('amount', ''),
                '{{DESCRIPTION}}': data.get('description', ''),
                '{{URL}}': data.get('url', data.get('pageUrl', '')),
                '{{TIMESTAMP}}': datetime.now().strftime('%Y-%m-%d %H:%M:%S')
            })
            
            # Add any custom fields from the data
            for key, value in data.items():
                if key not in ['title', 'date', 'amount', 'description', 'url', 'settings']:
                    str_value = self.process_array_value(value, array_handling)
                    replacements[f'{{{{{key.upper()}}}}}'] = str_value
        
        logger.info(f"Final replacements dictionary: {replacements}")  # Debug logging
        
        # Replace in paragraphs (preserving formatting)
        for paragraph in doc.paragraphs:
            self.replace_in_paragraph(paragraph, replacements)
        
        # Replace in tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for placeholder, replacement in replacements.items():
                        if placeholder in cell.text:
                            cell.text = cell.text.replace(placeholder, replacement)
        
        # Replace in headers and footers (preserving formatting)
        for section in doc.sections:
            # Header
            if section.header:
                logger.info(f"Processing header section...")
                for paragraph in section.header.paragraphs:
                    self.replace_in_paragraph(paragraph, replacements)
            
            # Footer
            if section.footer:
                logger.info(f"Processing footer section...")
                for paragraph in section.footer.paragraphs:
                    self.replace_in_paragraph(paragraph, replacements)
    
    def replace_in_paragraph(self, paragraph, replacements):
        """Replace placeholders in a paragraph while preserving formatting."""
        original_text = paragraph.text
        
        # Check if paragraph contains any placeholders
        has_placeholders = any(placeholder in original_text for placeholder in replacements.keys())
        if not has_placeholders:
            return
        
        logger.info(f"Processing paragraph with placeholders: '{original_text[:100]}...'")
        
        # Handle barcode placeholders first (they need special handling)
        barcode_placeholders = []
        for placeholder, replacement in replacements.items():
            if placeholder.startswith('{{BARCODE_') and placeholder in original_text:
                barcode_placeholders.append((placeholder, replacement))
                # Remove the barcode placeholder from text first
                self._replace_in_runs(paragraph, placeholder, "")
        
        # Process regular text replacements run by run
        for placeholder, replacement in replacements.items():
            if not placeholder.startswith('{{BARCODE_'):
                self._replace_in_runs(paragraph, placeholder, replacement)
        
        # Handle barcode placeholders after text replacements
        for placeholder, barcode_value in barcode_placeholders:
            logger.info(f"Inserting barcode for '{placeholder}' with value '{barcode_value}'")
            success = self.insert_barcode(paragraph, barcode_value)
            if success:
                logger.info(f"Successfully inserted barcode for {placeholder}")
            else:
                logger.warning(f"Failed to insert barcode for {placeholder}")
    
    def _replace_in_runs(self, paragraph, placeholder, replacement):
        """Replace placeholder in runs while preserving formatting."""
        if placeholder not in paragraph.text:
            return
            
        logger.info(f"Replacing '{placeholder}' with '{replacement}' in paragraph")
        
        # Find runs that contain part of the placeholder
        full_text = ""
        run_texts = []
        
        for run in paragraph.runs:
            run_texts.append(run.text)
            full_text += run.text
        
        # Find placeholder position in full text
        placeholder_start = full_text.find(placeholder)
        if placeholder_start == -1:
            return
            
        placeholder_end = placeholder_start + len(placeholder)
        
        # Find which runs contain the placeholder
        current_pos = 0
        affected_runs = []
        
        for i, run_text in enumerate(run_texts):
            run_start = current_pos
            run_end = current_pos + len(run_text)
            
            # Check if this run overlaps with placeholder
            if run_start < placeholder_end and run_end > placeholder_start:
                # Calculate overlap within this run
                overlap_start = max(0, placeholder_start - run_start)
                overlap_end = min(len(run_text), placeholder_end - run_start)
                affected_runs.append((i, overlap_start, overlap_end, run_text))
            
            current_pos = run_end
        
        # Replace text in affected runs
        if len(affected_runs) == 1:
            # Simple case: placeholder is entirely within one run
            run_index, start, end, run_text = affected_runs[0]
            new_text = run_text[:start] + replacement + run_text[end:]
            paragraph.runs[run_index].text = new_text
        else:
            # Complex case: placeholder spans multiple runs
            self._replace_across_runs(paragraph, affected_runs, placeholder, replacement)
    
    def _replace_across_runs(self, paragraph, affected_runs, placeholder, replacement):
        """Handle placeholder replacement across multiple runs."""
        if not affected_runs:
            return
        
        # Remove placeholder parts from all affected runs and put replacement in first run
        for i, (run_index, start, end, run_text) in enumerate(affected_runs):
            run = paragraph.runs[run_index]
            
            if i == 0:
                # First run: keep text before placeholder + replacement
                new_text = run_text[:start] + replacement
                run.text = new_text
            elif i == len(affected_runs) - 1:
                # Last run: keep text after placeholder
                new_text = run_text[end:]
                run.text = new_text
            else:
                # Middle runs: remove all text (it's part of placeholder)
                run.text = ""
    
    def insert_barcode(self, paragraph, barcode_value, barcode_type='code128'):
        """Insert a barcode into a paragraph (requires python-barcode)."""
        try:
            # Try to import barcode library
            from barcode import Code128, Code39, EAN13
            from barcode.writer import ImageWriter
            import io
            from docx.shared import Inches
            
            # Generate barcode image
            if barcode_type.lower() == 'code128':
                barcode_class = Code128
            elif barcode_type.lower() == 'code39':  
                barcode_class = Code39
            elif barcode_type.lower() == 'ean13':
                barcode_class = EAN13
            else:
                barcode_class = Code128  # Default
            
            # Create barcode
            barcode = barcode_class(str(barcode_value), writer=ImageWriter())
            
            # Save barcode to memory
            barcode_buffer = io.BytesIO()
            barcode.write(barcode_buffer)
            barcode_buffer.seek(0)
            
            # Insert into document
            run = paragraph.add_run()
            run.add_picture(barcode_buffer, width=Inches(1.2), height=Inches(0.4))
            
            logger.info(f"Inserted barcode for value: {barcode_value}")
            return True
            
        except ImportError:
            logger.warning("python-barcode not installed. Install with: pip install python-barcode[images]")
            # Fallback: just insert the text value
            paragraph.add_run(f"[BARCODE: {barcode_value}]")
            return False
        except Exception as e:
            logger.error(f"Error creating barcode: {e}")
            # Fallback: just insert the text value  
            paragraph.add_run(f"[BARCODE: {barcode_value}]")
            return False
    
    def get_field_value(self, data: Dict[str, Any], field_path: str, array_handling: str):
        """Extract value from data using field path (supports array indexing)."""
        try:
            # Handle array indexing like 'emails[0]'
            if '[' in field_path and ']' in field_path:
                field_name = field_path.split('[')[0]
                index_str = field_path.split('[')[1].split(']')[0]
                
                if field_name in data:
                    field_value = data[field_name]
                    if isinstance(field_value, list) and field_value:
                        try:
                            index = int(index_str)
                            if 0 <= index < len(field_value):
                                return field_value[index]
                        except ValueError:
                            pass
                return None
            
            # Handle simple field access
            if field_path in data:
                value = data[field_path]
                return self.process_array_value(value, array_handling)
            
            return None
            
        except Exception as e:
            logger.error(f"Error extracting field value '{field_path}': {e}")
            return None
    
    def process_array_value(self, value, array_handling: str):
        """Process array values according to handling preference."""
        if not isinstance(value, list):
            return str(value) if value is not None else ''
        
        if not value:  # Empty array
            return ''
        
        if array_handling == 'first':
            return str(value[0])
        elif array_handling == 'join_comma':
            return ', '.join(str(item) for item in value)
        elif array_handling == 'join_space':
            return ' '.join(str(item) for item in value)
        elif array_handling == 'join_newline':
            return '\n'.join(str(item) for item in value)
        elif array_handling == 'count':
            return str(len(value))
        else:
            # Default to first item
            return str(value[0])
    
    def apply_text_transform(self, text: str, transform: str) -> str:
        """Apply text transformation."""
        if not text or transform == 'none':
            return text
        
        if transform == 'uppercase':
            return text.upper()
        elif transform == 'lowercase':
            return text.lower()
        elif transform == 'capitalize':
            return text.title()
        elif transform == 'sentence':
            return text.capitalize()
        else:
            return text
    
    def replace_text_placeholders(self, content: str, data: Dict[str, Any]) -> str:
        """Replace placeholders in text content."""
        # Convert all data values to strings and create placeholder replacements
        replacements = {}
        
        logger.info(f"Raw extracted data keys: {list(data.keys())}")
        
        for key, value in data.items():
            placeholder = f"{{{{{key.upper()}}}}}"
            logger.info(f"Processing key: '{key}' -> placeholder: '{placeholder}'")
            
            if isinstance(value, list):
                # Handle arrays by joining with commas
                replacement = ', '.join(str(item) for item in value)
            elif isinstance(value, dict):
                # Handle objects by converting to string representation
                replacement = str(value)
            else:
                replacement = str(value) if value is not None else ''
            
            replacements[placeholder] = replacement
            logger.info(f"Will replace '{placeholder}' with '{replacement}'")
        
        logger.info(f"All replacements: {replacements}")
        
        # Replace all placeholders in the content
        processed_content = content
        for placeholder, replacement in replacements.items():
            if placeholder in content:
                logger.info(f"Found and replacing {placeholder}")
                processed_content = processed_content.replace(placeholder, replacement)
            else:
                logger.info(f"Placeholder {placeholder} not found in template")
        
        return processed_content
    
    def process_template(self, data: Dict[str, Any]) -> Dict[str, Any]:
        """Process a Word template with the provided data."""
        try:
            # Get template path - could be just a name or full path
            template_name = data.get('template', self.config['default_template'])
            
            # If it's just a filename, look in the default template directory
            if not Path(template_name).is_absolute():
                template_path = Path(self.config['template_path']) / template_name
            else:
                # If it's a full path, use it directly
                template_path = Path(template_name)
                
            # If still not found, search in all known template directories
            if not template_path.exists():
                # Try to find the template in any of our known directories
                template_dirs = [
                    Path(self.config['template_path']),
                    Path(__file__).parent.parent / 'templates'
                ]
                
                found = False
                for template_dir in template_dirs:
                    if template_dir.exists():
                        potential_path = template_dir / template_name
                        if potential_path.exists():
                            template_path = potential_path
                            found = True
                            break
                
                if not found:
                    return {
                        'success': False,
                        'error': f'Template not found: {template_name} (searched in {[str(d) for d in template_dirs]})'
                    }
            
            # Create output directory if it doesn't exist
            output_dir = Path(self.config['output_path'])
            output_dir.mkdir(parents=True, exist_ok=True)
            
            # Generate output filename with timestamp
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            base_name = template_path.stem
            
            # Handle different template types
            if template_path.suffix.lower() == '.txt':
                # Process text template
                output_filename = f"{base_name}_{timestamp}.txt"
                output_path = output_dir / output_filename
                
                logger.info(f"Processing text template: {template_path}")
                
                # Read the text template
                with open(template_path, 'r', encoding='utf-8') as f:
                    template_content = f.read()
                
                # Replace placeholders
                extracted_data = data.get('extractedData', {})
                logger.info(f"Extracted data for replacement: {extracted_data}")
                
                processed_content = self.replace_text_placeholders(template_content, extracted_data)
                
                # Save the processed text
                with open(output_path, 'w', encoding='utf-8') as f:
                    f.write(processed_content)
                    
                logger.info(f"Text document saved: {output_path}")
                
            else:
                # Process Word document template
                output_filename = f"{base_name}_{timestamp}.docx"
                output_path = output_dir / output_filename
                
                logger.info(f"Processing Word template: {template_path}")
                doc = Document(template_path)
                
                # Replace placeholders
                extracted_data = data.get('extractedData', {})
                logger.info(f"Extracted data for replacement: {extracted_data}")
                self.replace_placeholders(doc, extracted_data)
                
                # Save the updated document
                doc.save(output_path)
                logger.info(f"Word document saved: {output_path}")
            
            # Auto-open the document if configured
            if self.config.get('auto_open', True):
                try:
                    if os.name == 'nt':  # Windows
                        os.startfile(output_path)
                    elif os.name == 'posix':  # macOS and Linux
                        subprocess.run(['open' if sys.platform == 'darwin' else 'xdg-open', output_path])
                except Exception as e:
                    logger.warning(f"Could not auto-open document: {e}")
            
            return {
                'success': True,
                'output_path': str(output_path),
                'message': f'Document created successfully: {output_filename}'
            }
            
        except Exception as e:
            logger.error(f"Error processing template: {e}")
            return {
                'success': False,
                'error': str(e)
            }
    
    def handle_config_update(self, data: Dict[str, Any]) -> Dict[str, Any]:
        """Handle configuration updates."""
        try:
            # Update configuration
            for key in ['template_path', 'output_path', 'auto_open', 'default_template']:
                if key in data:
                    self.config[key] = data[key]
            
            self.save_config()
            
            return {
                'success': True,
                'message': 'Configuration updated successfully',
                'config': self.config
            }
            
        except Exception as e:
            logger.error(f"Error updating config: {e}")
            return {
                'success': False,
                'error': str(e)
            }
    
    def handle_get_config(self) -> Dict[str, Any]:
        """Handle configuration retrieval."""
        return {
            'success': True,
            'config': self.config
        }
    
    def handle_list_templates(self) -> Dict[str, Any]:
        """Handle template listing."""
        try:
            template_dirs = [
                Path(self.config['template_path']),  # Primary template directory
                Path(__file__).parent.parent / 'templates'  # Extension templates as fallback
            ]
            
            templates = []
            for template_dir in template_dirs:
                if template_dir.exists():
                    # Look for .docx files
                    for file_path in template_dir.glob('*.docx'):
                        if not file_path.name.startswith('~'):  # Skip temporary files
                            templates.append({
                                'name': file_path.name,
                                'path': str(file_path),
                                'size': file_path.stat().st_size,
                                'modified': file_path.stat().st_mtime
                            })
                    
                    # Also look for .txt files as simple templates (for testing)
                    for file_path in template_dir.glob('*.txt'):
                        if not file_path.name.startswith('~') and 'template' in file_path.name.lower():
                            templates.append({
                                'name': file_path.name,
                                'path': str(file_path),
                                'size': file_path.stat().st_size,
                                'modified': file_path.stat().st_mtime,
                                'type': 'text'
                            })
            
            if not templates:
                return {
                    'success': True,
                    'templates': [],
                    'message': f'No templates found in {[str(d) for d in template_dirs]}'
                }
            
            return {
                'success': True,
                'templates': templates
            }
            
        except Exception as e:
            logger.error(f"Error listing templates: {e}")
            return {
                'success': False,
                'error': str(e)
            }
    
    def run(self):
        """Main message processing loop."""
        logger.info("Word Template Updater started")
        
        while True:
            try:
                message = self.read_message()
                if message is None:
                    break
                
                logger.info(f"Received message: {message.get('action', 'unknown')}")
                
                action = message.get('action')
                response = {'success': False, 'error': 'Unknown action'}
                
                if action == 'update_template':
                    response = self.process_template(message.get('data', {}))
                elif action == 'update_config':
                    response = self.handle_config_update(message.get('data', {}))
                elif action == 'get_config':
                    response = self.handle_get_config()
                elif action == 'list_templates':
                    response = self.handle_list_templates()
                elif action == 'ping':
                    response = {'success': True, 'message': 'pong'}
                else:
                    response = {'success': False, 'error': f'Unknown action: {action}'}
                
                self.send_message(response)
                
            except KeyboardInterrupt:
                logger.info("Received interrupt signal")
                break
            except Exception as e:
                logger.error(f"Unexpected error: {e}")
                self.send_message({
                    'success': False,
                    'error': f'Unexpected error: {str(e)}'
                })
        
        logger.info("Word Template Updater stopped")

def main():
    """Main entry point."""
    try:
        updater = WordTemplateUpdater()
        updater.run()
    except Exception as e:
        logger.error(f"Fatal error: {e}")
        sys.exit(1)

if __name__ == '__main__':
    main()